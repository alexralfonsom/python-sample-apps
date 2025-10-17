#!/usr/bin/env python3
import re
import subprocess
from pathlib import Path
from docx import Document
import argparse


def extract_cc(docx_path: Path):
    """Extrae el número de cédula (CC) del documento Word.
    Prioriza el caso en el que 'CC' está en una celda y el valor en la celda adyacente de la misma fila.
    Si no se encuentra en tablas, realiza una búsqueda inline por regex.
    """
    # Valor numérico (6 a 20 dígitos)
    pattern_value = re.compile(r"([0-9]{6,20})")
    # Patrón inline alterno tipo 'CC: 123456789'
    pattern_inline = re.compile(r"\bCC\b\s*[:\-]?\s*([0-9]{6,20})", re.IGNORECASE)

    doc = Document(docx_path)

    # 1) Tablas: buscar filas donde una celda sea 'CC' y la celda adyacente tenga el número
    for table in doc.tables:
        for row in table.rows:
            cells = list(row.cells)
            for idx, cell in enumerate(cells):
                text = (cell.text or "").strip()
                # ¿Esta celda es la etiqueta 'CC'?
                if re.fullmatch(r"CC", text, flags=re.IGNORECASE) or re.fullmatch(r"CC\s*[:\-]?", text, flags=re.IGNORECASE):
                    # Tomar la celda inmediata a la derecha si existe
                    if idx + 1 < len(cells):
                        next_text = (cells[idx + 1].text or "").strip()
                        m_val = pattern_value.search(next_text)
                        if m_val:
                            return m_val.group(1)
                # Caso mixto: la misma celda tiene 'CC: 123...'
                m_inline_cell = pattern_inline.search(text)
                if m_inline_cell:
                    return m_inline_cell.group(1)

            # Fallback por fila: si alguna celda es 'CC' y otra tiene el número
            has_cc_label = any(re.fullmatch(r"\s*CC\s*[:\-]?\s*", (c.text or "").strip(), flags=re.IGNORECASE) for c in cells)
            if has_cc_label:
                for c in cells:
                    t = (c.text or "").strip()
                    # Evitar capturar otra vez la celda label
                    if not re.fullmatch(r"\s*CC\s*[:\-]?\s*", t, flags=re.IGNORECASE):
                        m_val = pattern_value.search(t)
                        if m_val:
                            return m_val.group(1)

    # 2) Párrafos: búsqueda inline (por si el formato cambia)
    for p in doc.paragraphs:
        m = pattern_inline.search(p.text or "")
        if m:
            return m.group(1)

    return None


def convert_docx_to_pdf(soffice_path: str, docx_path: Path, outdir: Path):
    """Convierte un archivo .docx a PDF usando LibreOffice."""
    cmd = [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"Error en LibreOffice: {result.stderr}")
    pdf_path = outdir / (docx_path.stem + ".pdf")
    return pdf_path


def main():
    parser = argparse.ArgumentParser(description="Convierte archivos DOCX a PDF usando CC como nombre.")
    parser.add_argument("--input", "-i", required=True, help="Carpeta con archivos DOCX")
    parser.add_argument("--output", "-o", required=True, help="Carpeta donde guardar los PDFs")
    parser.add_argument("--soffice", default="soffice",
                        help="Ruta a soffice (LibreOffice). En Windows suele ser 'C:\\Program Files\\LibreOffice\\program\\soffice.exe'")
    args = parser.parse_args()

    input_dir = Path(args.input).resolve()
    output_dir = Path(args.output).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    for docx_file in input_dir.glob("*.docx"):
        print(f"Procesando: {docx_file.name}")
        try:
            cc = extract_cc(docx_file)
            if not cc:
                print(f"⚠️ No se encontró CC en {docx_file.name}, usando nombre base.")
                cc = docx_file.stem

            pdf_path = convert_docx_to_pdf(args.soffice, docx_file, output_dir)
            new_name = output_dir / f"MB-{cc}.pdf"
            pdf_path.rename(new_name)
            print(f"✅ Generado: {new_name.name}")

        except Exception as e:
            print(f"❌ Error en {docx_file.name}: {e}")


if __name__ == "__main__":
    main()